"""MD -> DOCX via pandoc, then post-process fonts / sizes / heading colors.

rFonts behavior by primary:
  zh-primary: ascii = hAnsi = cs = eastAsia = cjk name  (whole doc in CJK font;
              CJK fonts contain Latin glyphs, so mixed text stays visually uniform)
  en-primary: ascii = hAnsi = cs = latin name; eastAsia = cjk name  (classic Word mix)
"""
from __future__ import annotations

import subprocess
from pathlib import Path

from . import themes


def _apply_style(out: Path, style: dict) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return  # python-docx not installed — keep pandoc defaults

    font_id = style.get("font", themes.DEFAULTS["font"])
    size_id = style.get("size", themes.DEFAULTS["size"])
    color_id = style.get("color", themes.DEFAULTS["color"])
    primary = style.get("primary", "zh")

    s = themes.get_size(size_id)
    c = themes.get_color(color_id)
    body_pt = float(s["body_pt"])  # type: ignore[arg-type]
    head_color = RGBColor.from_string(c["heading"])
    accent_dark = RGBColor.from_string(c["accent_dark"])

    def name_pair(role: str) -> tuple[str, str]:
        return themes.latin_name(font_id, role), themes.cjk_name(font_id, role)

    display_latin, display_cjk = name_pair("display")
    body_latin, body_cjk = name_pair("body")
    mono_latin, mono_cjk = name_pair("mono")

    doc = Document(str(out))

    def set_font(run, latin: str, cjk: str,
                 size_pt: float | None = None, color: RGBColor | None = None):
        # if zh-primary, use cjk everywhere (so even ASCII runs render in CJK font)
        ascii_name = cjk if primary == "zh" else latin
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"),    ascii_name)
        rFonts.set(qn("w:hAnsi"),    ascii_name)
        rFonts.set(qn("w:cs"),       ascii_name)
        rFonts.set(qn("w:eastAsia"), cjk)
        run.font.name = ascii_name
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if color is not None:
            run.font.color.rgb = color

    heading_scales = {
        "Heading 1": (body_pt * 2.0, head_color),
        "Heading 2": (body_pt * 1.5, accent_dark),
        "Heading 3": (body_pt * 1.2, head_color),
        "Heading 4": (body_pt * 1.08, accent_dark),
    }

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        if style_name in heading_scales:
            pt, color = heading_scales[style_name]
            for run in para.runs:
                set_font(run, display_latin, display_cjk, pt, color)
        else:
            for run in para.runs:
                is_code = bool(run.font.name) and "mono" in (run.font.name or "").lower()
                if is_code:
                    set_font(run, mono_latin, mono_cjk, body_pt)
                else:
                    set_font(run, body_latin, body_cjk, body_pt)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_font(run, body_latin, body_cjk, body_pt)

    doc.save(str(out))


def convert(src: Path, out: Path, assets_dir: Path, style: dict) -> None:
    cmd = ["pandoc", str(src), "-o", str(out), "--from", "gfm+smart", "--to", "docx"]
    ref = assets_dir / "reference.docx"
    if ref.exists():
        cmd += ["--reference-doc", str(ref)]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or "pandoc 失败")

    _apply_style(out, style)
